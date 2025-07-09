import os
import sys
import time
import logging
from pathlib import Path
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import win32serviceutil
import win32service
import win32event
import servicemanager
import socket
import re

class HTMLToDOCXConverter:
    """Converts HTML files to DOCX format while preserving formatting."""
    
    def __init__(self):
        self.downloads_path = self._get_downloads_path()
        self.setup_logging()
        
    def _get_downloads_path(self):
        """Get the Downloads folder path."""
        return str(Path.home() / "Downloads")
    
    def setup_logging(self):
        """Setup logging configuration."""
        log_dir = Path("C:/ProgramData/HTMLConverter/logs")
        log_dir.mkdir(parents=True, exist_ok=True)
        
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s',
            handlers=[
                logging.FileHandler(log_dir / "html_converter.log"),
                logging.StreamHandler()
            ]
        )
        self.logger = logging.getLogger(__name__)
    
    def convert_html_to_docx(self, html_file_path):
        """Convert HTML file to DOCX format."""
        try:
            html_path = Path(html_file_path)
            docx_path = html_path.with_suffix('.docx')
            
            # Read HTML file
            with open(html_path, 'r', encoding='utf-8') as file:
                html_content = file.read()
            
            # Parse HTML
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Create Word document
            doc = Document()
            
            # Set minimal margins (0.5cm = 0.2 inches)
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.2)
                section.bottom_margin = Inches(0.2)
                section.left_margin = Inches(0.2)
                section.right_margin = Inches(0.2)
            
            # Extract and apply CSS styles
            css_styles = self._extract_css_styles(soup)
            
            # Extract title
            title = soup.find('title')
            if title:
                doc.add_heading(title.get_text(), 0)
            
            # Process body content
            body = soup.find('body')
            if body:
                self._process_html_elements(body, doc, css_styles)
            else:
                # If no body tag, process the entire HTML
                self._process_html_elements(soup, doc, css_styles)
            
            # Save DOCX file
            doc.save(docx_path)
            self.logger.info(f"Successfully converted {html_path.name} to {docx_path.name}")
            
            # Remove original HTML file
            html_path.unlink()
            self.logger.info(f"Removed original HTML file: {html_path.name}")
            
            return True
            
        except Exception as e:
            self.logger.error(f"Error converting {html_file_path}: {str(e)}")
            return False
    
    def _extract_css_styles(self, soup):
        """Extract CSS styles from HTML."""
        css_styles = {}
        
        # Extract inline styles
        for tag in soup.find_all(style=True):
            if tag.get('style'):
                css_styles[tag] = self._parse_inline_css(tag.get('style'))
        
        # Extract style tags
        for style_tag in soup.find_all('style'):
            if style_tag.string:
                css_styles.update(self._parse_css_rules(style_tag.string))
        
        return css_styles
    
    def _parse_inline_css(self, style_string):
        """Parse inline CSS styles."""
        styles = {}
        if style_string:
            for rule in style_string.split(';'):
                if ':' in rule:
                    property_name, value = rule.split(':', 1)
                    styles[property_name.strip()] = value.strip()
        return styles
    
    def _parse_css_rules(self, css_text):
        """Parse CSS rules from style tags."""
        styles = {}
        # Simple CSS parser for basic rules
        rules = re.findall(r'([^{]+)\{([^}]+)\}', css_text)
        for selector, declarations in rules:
            selector = selector.strip()
            style_dict = {}
            for decl in declarations.split(';'):
                if ':' in decl:
                    prop, value = decl.split(':', 1)
                    style_dict[prop.strip()] = value.strip()
            styles[selector] = style_dict
        return styles
    
    def _apply_css_styles(self, element, paragraph, css_styles):
        """Apply CSS styles to Word document elements."""
        try:
            # Get inline styles
            inline_styles = {}
            if element.get('style'):
                inline_styles = self._parse_inline_css(element.get('style'))
            
            # Apply text formatting
            for run in paragraph.runs:
                # Font size
                if 'font-size' in inline_styles:
                    size_str = inline_styles['font-size']
                    if 'px' in size_str:
                        size = int(size_str.replace('px', ''))
                        run.font.size = Pt(size)
                
                # Font weight (bold)
                if 'font-weight' in inline_styles:
                    weight = inline_styles['font-weight']
                    if weight in ['bold', '700', '800', '900']:
                        run.font.bold = True
                
                # Font style (italic)
                if 'font-style' in inline_styles:
                    style = inline_styles['font-style']
                    if style == 'italic':
                        run.font.italic = True
                
                # Text color
                if 'color' in inline_styles:
                    color = inline_styles['color']
                    if color.startswith('#'):
                        # Convert hex to RGB
                        hex_color = color.lstrip('#')
                        rgb = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
                        run.font.color.rgb = RGBColor(*rgb)
                
                # Text decoration (underline)
                if 'text-decoration' in inline_styles:
                    decoration = inline_styles['text-decoration']
                    if 'underline' in decoration:
                        run.font.underline = True
            
            # Apply paragraph formatting
            # Text alignment
            if 'text-align' in inline_styles:
                align = inline_styles['text-align']
                if align == 'center':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif align == 'right':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif align == 'justify':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Line height
            if 'line-height' in inline_styles:
                line_height = inline_styles['line-height']
                if line_height.endswith('px'):
                    height = int(line_height.replace('px', ''))
                    paragraph.paragraph_format.line_spacing = Pt(height)
            
        except Exception as e:
            self.logger.warning(f"Error applying CSS styles: {str(e)}")
    
    def _process_html_elements(self, element, doc, css_styles):
        """Recursively process HTML elements and add them to the Word document."""
        for child in element.children:
            if child.name is None:  # Text node
                if child.strip():
                    paragraph = doc.add_paragraph(child.strip())
                    self._apply_css_styles(element, paragraph, css_styles)
            elif child.name == 'h1':
                heading = doc.add_heading(child.get_text(), level=1)
                self._apply_css_styles(child, heading, css_styles)
            elif child.name == 'h2':
                heading = doc.add_heading(child.get_text(), level=2)
                self._apply_css_styles(child, heading, css_styles)
            elif child.name == 'h3':
                heading = doc.add_heading(child.get_text(), level=3)
                self._apply_css_styles(child, heading, css_styles)
            elif child.name == 'h4':
                heading = doc.add_heading(child.get_text(), level=4)
                self._apply_css_styles(child, heading, css_styles)
            elif child.name == 'h5':
                heading = doc.add_heading(child.get_text(), level=5)
                self._apply_css_styles(child, heading, css_styles)
            elif child.name == 'h6':
                heading = doc.add_heading(child.get_text(), level=6)
                self._apply_css_styles(child, heading, css_styles)
            elif child.name == 'p':
                paragraph = doc.add_paragraph()
                # Handle mixed content (text and inline elements)
                self._process_mixed_content(child, paragraph, css_styles)
                self._apply_css_styles(child, paragraph, css_styles)
            elif child.name == 'br':
                doc.add_paragraph()
            elif child.name == 'ul':
                for li in child.find_all('li', recursive=False):
                    paragraph = doc.add_paragraph(li.get_text(), style='List Bullet')
                    self._apply_css_styles(li, paragraph, css_styles)
            elif child.name == 'ol':
                for li in child.find_all('li', recursive=False):
                    paragraph = doc.add_paragraph(li.get_text(), style='List Number')
                    self._apply_css_styles(li, paragraph, css_styles)
            elif child.name == 'table':
                self._process_table(child, doc, css_styles)
            elif child.name in ['div', 'span', 'section', 'article']:
                # Recursively process nested elements
                self._process_html_elements(child, doc, css_styles)
    
    def _process_mixed_content(self, element, paragraph, css_styles):
        """Process elements with mixed content (text and inline elements)."""
        for content in element.contents:
            if content.name is None:  # Text node
                if content.strip():
                    run = paragraph.add_run(content.strip())
                    self._apply_css_styles(element, paragraph, css_styles)
            elif content.name in ['strong', 'b']:
                run = paragraph.add_run(content.get_text())
                run.font.bold = True
            elif content.name in ['em', 'i']:
                run = paragraph.add_run(content.get_text())
                run.font.italic = True
            elif content.name == 'u':
                run = paragraph.add_run(content.get_text())
                run.font.underline = True
            elif content.name == 'span':
                run = paragraph.add_run(content.get_text())
                # Apply span-specific styles
                if content.get('style'):
                    span_styles = self._parse_inline_css(content.get('style'))
                    if 'color' in span_styles:
                        color = span_styles['color']
                        if color.startswith('#'):
                            hex_color = color.lstrip('#')
                            rgb = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
                            run.font.color.rgb = RGBColor(*rgb)
            else:
                # For other elements, just add the text
                run = paragraph.add_run(content.get_text())
    
    def _process_table(self, table_element, doc, css_styles):
        """Process HTML table and add it to the Word document."""
        rows = table_element.find_all('tr')
        if not rows:
            return
        
        # Create table
        table = doc.add_table(rows=len(rows), cols=0)
        table.style = 'Table Grid'
        
        for i, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            if i == 0:  # First row determines column count
                table.columns = len(cells)
            
            for j, cell in enumerate(cells):
                if j < len(table.rows[i].cells):
                    table_cell = table.rows[i].cells[j]
                    table_cell.text = cell.get_text().strip()
                    
                    # Apply cell-specific styles
                    if cell.get('style'):
                        cell_styles = self._parse_inline_css(cell.get('style'))
                        # Apply background color if specified
                        if 'background-color' in cell_styles:
                            bg_color = cell_styles['background-color']
                            if bg_color.startswith('#'):
                                hex_color = bg_color.lstrip('#')
                                rgb = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
                                # Note: Background color for table cells requires more complex handling
                                # This is a basic implementation


class DownloadFolderHandler(FileSystemEventHandler):
    """Handles file system events in the Downloads folder."""
    
    def __init__(self, converter):
        self.converter = converter
        self.logger = converter.logger
    
    def on_created(self, event):
        """Handle file creation events."""
        if not event.is_directory:
            file_path = Path(event.src_path)
            if file_path.suffix.lower() == '.html':
                self.logger.info(f"New HTML file detected: {file_path.name}")
                # Wait a moment to ensure file is fully written
                time.sleep(1)
                self.converter.convert_html_to_docx(str(file_path))


class HTMLConverterService(win32serviceutil.ServiceFramework):
    """Windows Service for HTML to DOCX conversion."""
    
    _svc_name_ = "HTMLToDOCXConverter"
    _svc_display_name_ = "HTML to DOCX Converter Service"
    _svc_description_ = "Monitors Downloads folder and converts HTML files to DOCX format"
    
    def __init__(self, args):
        win32serviceutil.ServiceFramework.__init__(self, args)
        self.stop_event = win32event.CreateEvent(None, 0, 0, None)
        self.converter = HTMLToDOCXConverter()
        self.observer = None
    
    def SvcStop(self):
        """Stop the service."""
        self.logger.info("Stopping HTML to DOCX Converter Service...")
        self.ReportServiceStatus(win32service.SERVICE_STOP_PENDING)
        win32event.SetEvent(self.stop_event)
        if self.observer:
            self.observer.stop()
    
    def SvcDoRun(self):
        """Run the service."""
        self.logger.info("Starting HTML to DOCX Converter Service...")
        self.main()
    
    def main(self):
        """Main service loop."""
        try:
            # Create event handler and observer
            event_handler = DownloadFolderHandler(self.converter)
            self.observer = Observer()
            self.observer.schedule(event_handler, self.converter.downloads_path, recursive=False)
            self.observer.start()
            
            self.logger.info(f"Monitoring Downloads folder: {self.converter.downloads_path}")
            
            # Keep service running
            while True:
                # Check if service should stop
                if win32event.WaitForSingleObject(self.stop_event, 1000) == win32event.WAIT_OBJECT_0:
                    break
                
        except Exception as e:
            self.logger.error(f"Service error: {str(e)}")
        finally:
            if self.observer:
                self.observer.stop()
                self.observer.join()


def run_as_console():
    """Run the converter as a console application for testing."""
    converter = HTMLToDOCXConverter()
    event_handler = DownloadFolderHandler(converter)
    observer = Observer()
    observer.schedule(event_handler, converter.downloads_path, recursive=False)
    observer.start()
    
    print(f"Monitoring Downloads folder: {converter.downloads_path}")
    print("Press Ctrl+C to stop...")
    
    try:
        while True:
            time.sleep(1)
    except KeyboardInterrupt:
        observer.stop()
        observer.join()
        print("Monitoring stopped.")


if __name__ == '__main__':
    if len(sys.argv) == 1:
        # Run as console application
        run_as_console()
    elif len(sys.argv) > 1 and sys.argv[1] == 'test':
        # Run test mode for standalone executable
        print("HTML to DOCX Converter - Test Mode")
        print("=" * 40)
        
        # Basic test without external dependencies
        converter = HTMLToDOCXConverter()
        test_html = """
<!DOCTYPE html>
<html>
<head><title>Standalone Test</title></head>
<body>
<h1>Test Document</h1>
<p>This is a test paragraph with <strong>bold</strong> and <em>italic</em> text.</p>
<ul>
    <li>List item 1</li>
    <li>List item 2</li>
</ul>
<table border="1">
    <tr><th>Header</th><th>Value</th></tr>
    <tr><td>Test</td><td>Success</td></tr>
</table>
</body>
</html>
        """
        downloads_path = Path(converter.downloads_path)
        test_file = downloads_path / "standalone_test.html"
        
        try:
            print("Creating test HTML file...")
            with open(test_file, 'w', encoding='utf-8') as f:
                f.write(test_html.strip())
            
            print("Converting HTML to DOCX...")
            success = converter.convert_html_to_docx(str(test_file))
            
            if success:
                docx_path = test_file.with_suffix('.docx')
                if docx_path.exists():
                    print("✓ Test completed successfully!")
                    print(f"✓ DOCX file created: {docx_path.name}")
                    print("✓ Original HTML file removed")
                    print("\nThe converter is working correctly!")
                else:
                    print("✗ Test failed: DOCX file not found")
            else:
                print("✗ Test failed: Conversion unsuccessful")
                
        except Exception as e:
            print(f"✗ Test error: {e}")
        
        print("\nPress Enter to exit...")
        input()
    else:
        # Run as Windows service
        win32serviceutil.HandleCommandLine(HTMLConverterService) 